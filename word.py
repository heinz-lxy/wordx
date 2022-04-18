from docxcompose.composer import Composer
from docx import Document
from template import Template
from fake_zip import FakeZip
from lxml import etree
from lxml.builder import E
from io import BytesIO
from zipfile import ZipFile
import t 
import random


def make_res_id():
    """生成资源id"""
    return 'rId' + str(random.randint(1000,9999))  # rId9896


class Word(object):
    """word文档生成"""
    def __init__(self):
        self.engine = Template()
        self.file = ''
        self.data = []

    def add_rel(self, src, res_id, res_type, filename):
        """添加映射"""
        tree = etree.parse(BytesIO(src))
        root = tree.getroot()
        res_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/%s' % res_type
        rel_element = E.Relationship(Id=res_id, Type=res_type, Target=filename)
        root.append(rel_element)
        return etree.tostring(root)

    def add_element(self, ele_type, content):
        """添加元素"""
        self.data.append({
            'type': ele_type,
            'content': content
        })

    def add_paragraph(self, text):
        """添加段落"""
        self.add_element('paragraph', text)

    def add_picture(self, path):
        """添加图片"""
        self.add_element('picture', path)

    def add_table(self, data):
        """添加表格"""
        self.add_element('table', data)

    def add_page_break(self):
        """另起一页"""
        self.add_element('blank', '')

    def set_horizontal(self):
        """设为横版"""
        self.add_element('horizontal', '')

    def set_vertical(self):
        """设为横版"""
        self.add_element('vertical', '')

    def add_header(self):
        """添加页眉"""
        self.add_element('header', '我是页眉') 

    def add_catalog(self, content):
        """添加目录"""
        self.add_element('catalog', content)

    def create_header(self):
        pass

    def create_footer(self):
        pass

    def res_preprocess(self):
        """资源文件处理"""
        src = 'templates/template.docx'
        dst = r'output\%s2.docx' % t.timestamp()
        zip1 = ZipFile(src, 'r') 
        zip2 = ZipFile(dst, 'x')
        tmp = []
        for fileinfo in zip1.infolist():
            content = zip1.open(fileinfo).read()
            if(fileinfo.filename == 'word/_rels/document.xml.rels'):
                for item in self.data:
                    if item['type'] == 'image':
                        pic_path = item['content']['path']
                        pic_data = open(pic_path, 'rb').read()
                        filename = pic_path.split('\\')[-1]  
                        filename = filename.split('.')[0] + '.jpeg'  # 图片文件扩展名必须为jpeg
                        res_id = make_res_id()
                        content = self.add_rel(content, res_id, 'image', 'media/%s' % filename)  # 添加映射
                        zip2.writestr('word/media/%s' % filename, pic_data)  # 文件复制
                        item['content']['res_id'] = res_id
                        print(item['content']['path'])
                    if item['type'] == 'header':   
                        header_content = self.engine.render_template('resource/header.xml', content=item['content']) 
                        res_id = make_res_id()
                        filename = 'header' + res_id[3:] + '.xml'
                        content = self.add_rel(content, res_id, 'header', '%s' % filename)  
                        zip2.writestr('word/%s' % filename, header_content)  
                        item['content']['res_id'] = res_id
                    tmp.append(item)
                zip2.writestr('word/_rels/document.xml.rels', content)
            else:
                zip2.writestr(fileinfo.filename, content)
        print(tmp)
        return dst, tmp

    def generate(self):
        """生成文档"""
        template_xml = 'template.xml'
        template_docx, data = self.res_preprocess()  # 经过预处理的docx文档
        dst = r'output\%s1.docx' % t.timestamp()
        xml_data = self.engine.render(template_xml, data)
        return Zip(template_docx).replace('word/document.xml', xml_data)

    @classmethod
    def merge(cls, doc_list):
        """合并多个docx文档"""
        dst = r'output\%s.docx' % t.timestamp()
        tmp = Document()
        composer = Composer(tmp)
        for doc in doc_list:
            doc = Document(doc)
            doc.add_page_break()
            composer.append(doc)    
        composer.save(dst)
        return dst

  
if __name__ == '__main__':
    word = Word()
    word.add_paragraph('haha')
    word.set_horizontal()
    # word.add_catalog('')
    # word.add_page_break()
    # word.add_header()

    # word.add_picture(r'C:\Users\mi\Desktop\store.jpeg')
    # word.add_table({
    #         'header': ['姓名','年龄'], 
    #         # 'width': [3079, 3079],
    #         'data': [[('中文名',0),('英文名',0),'化学文摘号',('OELs (mg/m3)',3),('临界不良健康效应',0),('备注',0)],['v','v','CAS 号','MAC','PC-TWA','PC-STEL','v','v']]
    #     }
    # )
    rst = word.generate()
    print(rst)
    t.open_file(rst)






